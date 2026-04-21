from __future__ import annotations

from io import BytesIO
from pathlib import Path

from pypdf import PdfReader
from docx import Document


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()

    if suffix == ".docx":
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs).strip()

    if suffix == ".doc":
        try:
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            if paragraphs:
                return "\n".join(paragraphs).strip()
        except Exception:
            pass
        return content.decode("utf-8", errors="ignore").strip()

    return content.decode("utf-8", errors="ignore").strip()


def chunk_text(text: str, chunk_size: int = 900, chunk_overlap: int = 180) -> list[str]:
    cleaned = "\n".join(line.strip() for line in text.splitlines()).strip()
    if not cleaned:
        return []

    if chunk_size <= 0:
        return [cleaned]

    chunks: list[str] = []
    start = 0
    length = len(cleaned)
    overlap = max(0, min(chunk_overlap, chunk_size - 1))

    while start < length:
        end = min(length, start + chunk_size)
        chunk = cleaned[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= length:
            break
        start = max(end - overlap, start + 1)

    return chunks
